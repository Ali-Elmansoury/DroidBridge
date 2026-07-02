"""Generate docs/DroidBridge_Project_Document.docx from the Markdown source."""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Install python-docx first: pip install python-docx")
    sys.exit(1)

ROOT = Path(__file__).parent.parent
SRC  = ROOT / "docs" / "DroidBridge_Project_Document.md"
DST  = ROOT / "docs" / "DroidBridge_Project_Document.docx"


def _get_app_version() -> str:
    """Read __version__ from droidbridge/__init__.py without importing the package."""
    init = ROOT / "droidbridge" / "__init__.py"
    m = re.search(r'^__version__\s*=\s*["\']([^"\']+)["\']', init.read_text(), re.MULTILINE)
    return m.group(1) if m else "1.0.0"


def _inject_version(md: str, version: str) -> str:
    """Replace the Version line in the markdown preamble with the current app version."""
    return re.sub(
        r"^(Version\s+)\d+\.\d+\.\d+",
        rf"\g<1>{version}",
        md,
        count=1,
        flags=re.MULTILINE,
    )

# ── Colours ───────────────────────────────────────────────────────────────────
BLUE_HDG  = RGBColor(0x2E, 0x75, 0xB6)   # section-heading blue
NAVY_HEX  = "1F497D"                       # dark-navy hex for banners/table headers
NAVY_RGB  = RGBColor(0x1F, 0x49, 0x7D)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GREY      = RGBColor(0x59, 0x59, 0x59)    # footer grey

# callout box (fill_hex, text_RGBColor)
BOX_BLUE   = ("DEEAF1", RGBColor(0x1F, 0x49, 0x7D))
BOX_GREEN  = ("E2EFDA", RGBColor(0x37, 0x56, 0x23))
BOX_ORANGE = ("FCE4D6", RGBColor(0xC5, 0x5A, 0x11))

CODE_BG = "1E2D3D"
CODE_FG = RGBColor(0xD4, 0xD4, 0xD4)

# Emoji pattern for stripping from headings
_EMOJI = re.compile(
    "[\U00010000-\U0010ffff"
    "☀-⛿✀-➿"
    "\U0001F300-\U0001F9FF]"
)


# ── XML helpers ───────────────────────────────────────────────────────────────

def _shade_para(p, fill_hex: str):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def _shade_cell(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _bottom_border(p, color="CCCCCC", sz="4"):
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:pBdr")):
        pPr.remove(old)
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), sz)
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def _field_runs(para, field_name: str):
    """Append PAGE / NUMPAGES field as three runs."""
    for ftype, text in [("begin", None), (None, field_name), ("end", None)]:
        r = para.add_run()
        r.font.size = Pt(9)
        r.font.color.rgb = GREY
        if ftype:
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), ftype)
            r._r.append(fc)
        else:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = f" {text} "
            r._r.append(instr)


# ── Inline formatting ─────────────────────────────────────────────────────────

def _inline(para, text: str, bold=False, color=None, size=None):
    """Add runs with **bold** and `code` inline markers honoured."""
    for tok in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = para.add_run(tok[2:-2])
            r.bold = True
            if color:
                r.font.color.rgb = color
            if size:
                r.font.size = Pt(size)
        elif tok.startswith("`") and tok.endswith("`"):
            r = para.add_run(tok[1:-1])
            r.font.name = "Courier New"
            r.font.size = Pt(9.5)
        else:
            r = para.add_run(tok)
            if bold:
                r.bold = True
            if color:
                r.font.color.rgb = color
            if size:
                r.font.size = Pt(size)


# ── Title page ────────────────────────────────────────────────────────────────

def _title_page(doc, title: str, subtitle: str, version: str):
    # top spacer
    for _ in range(10):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(2)

    # "DROIDBRIDGE"
    parts = title.split(" — ", 1)
    app_name = parts[0].upper()
    app_line = parts[1] if len(parts) > 1 else ""

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(app_name)
    r.bold = True
    r.font.size = Pt(52)
    r.font.color.rgb = BLUE_HDG
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)

    # tagline
    if app_line:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(app_line)
        r2.font.size = Pt(16)
        p2.paragraph_format.space_after = Pt(12)

    # blue decorative separator bar — centered, taller
    p_bar = doc.add_paragraph()
    p_bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bar.paragraph_format.left_indent = Inches(1.5)
    p_bar.paragraph_format.right_indent = Inches(1.5)
    p_bar.paragraph_format.space_before = Pt(4)
    p_bar.paragraph_format.space_after = Pt(14)
    _shade_para(p_bar, "2E75B6")
    r_bar = p_bar.add_run(" ")
    r_bar.font.size = Pt(18)

    # document-type subtitle (italic)
    clean = re.sub(r"\*+", "", subtitle).strip()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(clean)
    r3.italic = True
    r3.font.size = Pt(14)
    p3.paragraph_format.space_after = Pt(18)

    # version / project lines
    ver_parts = [v.strip() for v in version.split("|")]
    ver_main  = " | ".join(ver_parts[:2]) if len(ver_parts) >= 2 else version
    ver_extra = ver_parts[2] if len(ver_parts) >= 3 else ""

    for txt in filter(None, [ver_main, ver_extra]):
        pv = doc.add_paragraph()
        pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rv = pv.add_run(txt)
        rv.font.size = Pt(11)
        rv.font.color.rgb = GREY
        pv.paragraph_format.space_after = Pt(4)

    doc.add_page_break()


# ── TOC page ──────────────────────────────────────────────────────────────────

def _toc_page(doc):
    p_h = doc.add_heading("Table of Contents", level=2)
    p_h.paragraph_format.space_before = Pt(0)
    _bottom_border(p_h)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    for ftype, text in [("begin", None), (None, r'TOC \o "1-3" \h \z \u'), ("separate", None)]:
        r = p.add_run()
        if ftype:
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), ftype)
            r._r.append(fc)
        else:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = f" {text} "
            r._r.append(instr)

    p.add_run("[Right-click → Update Field to generate table of contents]")

    r_end = p.add_run()
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    r_end._r.append(fc_end)

    doc.add_page_break()


# ── Footer ────────────────────────────────────────────────────────────────────

def _setup_footer(doc):
    for section in doc.sections:
        fp = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run("DroidBridge Project Document  |  Page ")
        r.font.size = Pt(9)
        r.font.color.rgb = GREY
        _field_runs(fp, "PAGE")
        r2 = fp.add_run("  of  ")
        r2.font.size = Pt(9)
        r2.font.color.rgb = GREY
        _field_runs(fp, "NUMPAGES")


# ── Banner heading (Module / Phase) ───────────────────────────────────────────

def _banner(doc, text: str):
    """Dark-navy filled box, white centered bold — used for Module/Phase headings."""
    p = doc.add_heading(text, level=2)   # keep in TOC via Heading 2 style
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(14)
    _shade_para(p, NAVY_HEX)
    for r in p.runs:
        r.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = WHITE
        r.font.name = "Calibri"


def _is_banner(text: str) -> bool:
    """True when the heading (after stripping leading number) is a Module/Phase banner."""
    t = text.lstrip("0123456789. ").strip().lower()
    return t.startswith("module") or t.startswith("phase")


# ── Callout boxes ─────────────────────────────────────────────────────────────

def _box_style(raw: str):
    t = re.sub(r"\*+", "", raw).strip().lower()
    if t.startswith("important"):
        return BOX_ORANGE
    if t.startswith(("note:", "note ", "recommendation", "recommended")):
        return BOX_GREEN
    return BOX_BLUE


def _callout(doc, bq_lines: list):
    text = " ".join(re.sub(r"^>\s*", "", l) for l in bq_lines).strip()
    fill, color = _box_style(text)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.right_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    _shade_para(p, fill)
    _inline(p, text, color=color, size=11)
    for r in p.runs:
        r.font.color.rgb = color
        if r.font.size is None:
            r.font.size = Pt(11)


# ── Tables ────────────────────────────────────────────────────────────────────

def _flush_table(doc, tbl_lines: list):
    rows = [r for r in tbl_lines if not re.match(r"^\|[-| :]+\|$", r.strip())]
    if not rows:
        return
    cols = [c.strip() for c in rows[0].strip("|").split("|")]
    data = []
    for row in rows[1:]:
        cells = [c.strip() for c in row.strip("|").split("|")]
        while len(cells) < len(cols):
            cells.append("")
        data.append(cells[: len(cols)])

    t = doc.add_table(rows=1 + len(data), cols=len(cols))
    t.style = "Table Grid"

    # header row
    for ci, h in enumerate(cols):
        cell = t.rows[0].cells[ci]
        _shade_cell(cell, NAVY_HEX)
        cell.paragraphs[0].clear()
        _inline(cell.paragraphs[0], h, bold=True, color=WHITE, size=11)
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = WHITE

    # data rows
    for ri, row in enumerate(data):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.paragraphs[0].clear()
            _inline(cell.paragraphs[0], val)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── Main builder ──────────────────────────────────────────────────────────────

def build_docx(md_text: str) -> Document:
    doc = Document()

    # page margins
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    # Normal style
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # Heading styles
    for lvl, sz, color in [
        (2, 20, BLUE_HDG),
        (3, 14, BLUE_HDG),
        (4, 12, RGBColor(0, 0, 0)),
    ]:
        h = doc.styles[f"Heading {lvl}"]
        h.font.name  = "Calibri"
        h.font.bold  = True
        h.font.size  = Pt(sz)
        h.font.color.rgb = color
        h.paragraph_format.space_before    = Pt(14 if lvl == 2 else 10)
        h.paragraph_format.space_after     = Pt(6  if lvl == 2 else 4)
        h.paragraph_format.keep_with_next  = True

    lines = md_text.split("\n")
    n = len(lines)
    i = 0

    in_code       = False
    code_lines    = []
    tbl_lines     = []
    bq_lines      = []
    in_banner_sec = False   # inside a Module/Phase heading?
    in_preamble   = True    # still on the title-page preamble?
    title_text    = None
    subtitle_text = ""
    version_text  = ""

    def flush_bq():
        nonlocal bq_lines
        if bq_lines:
            _callout(doc, bq_lines)
            bq_lines = []

    def flush_tbl():
        nonlocal tbl_lines
        if tbl_lines:
            _flush_table(doc, tbl_lines)
            tbl_lines = []

    while i < n:
        line = lines[i]

        # ── fenced code block ──────────────────────────────────────────
        if line.strip().startswith("```"):
            flush_bq(); flush_tbl()
            if not in_code:
                in_code = True; code_lines = []
            else:
                in_code = False
                if code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after  = Pt(4)
                    p.paragraph_format.left_indent  = Inches(0.2)
                    p.paragraph_format.right_indent = Inches(0.2)
                    _shade_para(p, CODE_BG)
                    r = p.add_run("\n".join(code_lines))
                    r.font.name = "Courier New"
                    r.font.size = Pt(8.5)
                    r.font.color.rgb = CODE_FG
                code_lines = []
            i += 1; continue

        if in_code:
            code_lines.append(line); i += 1; continue

        # ── table ──────────────────────────────────────────────────────
        if line.strip().startswith("|"):
            flush_bq()
            tbl_lines.append(line); i += 1; continue
        else:
            flush_tbl()

        # ── blockquote ─────────────────────────────────────────────────
        if line.strip().startswith(">"):
            if not in_preamble:
                bq_lines.append(line)
            i += 1; continue
        else:
            flush_bq()

        # ── horizontal rule ────────────────────────────────────────────
        if re.match(r"^---+\s*$", line.strip()):
            if in_preamble and title_text:
                _title_page(doc, title_text, subtitle_text, version_text)
                _setup_footer(doc)
                _toc_page(doc)
                in_preamble = False
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after  = Pt(4)
                _bottom_border(p, "CCCCCC", "6")
            i += 1; continue

        # ── headings ───────────────────────────────────────────────────
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            lvl  = len(m.group(1))
            text = _EMOJI.sub("", m.group(2)).strip()

            # H1 — title page info collection
            if lvl == 1:
                title_text = text
                i += 1
                # subtitle (bold line)
                while i < n and not lines[i].strip():
                    i += 1
                if i < n and lines[i].strip().startswith("**"):
                    subtitle_text = lines[i].strip(); i += 1
                # version line
                while i < n and not lines[i].strip():
                    i += 1
                if i < n and not lines[i].startswith(("#", ">", "-")):
                    version_text = lines[i].strip(); i += 1
                continue

            # H2
            if lvl == 2:
                if _is_banner(text):
                    _banner(doc, text)
                    in_banner_sec = True
                else:
                    p = doc.add_heading(text, level=2)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    _bottom_border(p, "CCCCCC", "4")
                    in_banner_sec = False
                i += 1; continue

            # H3
            if lvl == 3:
                if _is_banner(text):
                    _banner(doc, text)
                    in_banner_sec = True
                elif in_banner_sec:
                    # module/phase subsection — black bold, not blue
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(10)
                    p.paragraph_format.space_after  = Pt(4)
                    r = p.add_run(text)
                    r.bold = True
                    r.font.size = Pt(13)
                    r.font.name = "Calibri"
                else:
                    doc.add_heading(text, level=3)
                i += 1; continue

            # H4
            if lvl == 4:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after  = Pt(4)
                r = p.add_run(text)
                r.bold = True
                r.font.size = Pt(12)
                r.font.name = "Calibri"
                i += 1; continue

            i += 1; continue

        # ── ordered list ───────────────────────────────────────────────
        m = re.match(r"^\d+\.\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(2)
            _inline(p, m.group(1))
            i += 1; continue

        # ── unordered list ─────────────────────────────────────────────
        m = re.match(r"^(\s*)[*\-]\s+(.*)", line)
        if m:
            indent = len(m.group(1))
            try:
                style = "List Bullet 2" if indent >= 2 else "List Bullet"
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25 + 0.25 * (indent >= 2))
            p.paragraph_format.space_after = Pt(2)
            _inline(p, m.group(2))
            i += 1; continue

        # ── empty line ─────────────────────────────────────────────────
        if not line.strip():
            i += 1; continue

        # ── normal paragraph ───────────────────────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _inline(p, line.strip())
        i += 1

    flush_bq(); flush_tbl()
    return doc


if __name__ == "__main__":
    version = _get_app_version()
    print(f"Reading  {SRC}")
    md = SRC.read_text(encoding="utf-8")
    md = _inject_version(md, version)
    SRC.write_text(md, encoding="utf-8")   # keep .md in sync
    print("Building docx…")
    doc = build_docx(md)
    doc.save(DST)
    print(f"Written  {DST}  ({DST.stat().st_size // 1024} KB)")
